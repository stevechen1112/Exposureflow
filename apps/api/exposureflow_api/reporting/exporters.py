"""Report export helpers — Markdown, PDF, DOCX."""

from __future__ import annotations

import io
from typing import Literal

ExportFormat = Literal["markdown", "pdf", "docx"]


def export_markdown(content: str) -> bytes:
    return content.encode("utf-8")


def export_pdf(content: str, *, title: str = "ExposureFlow Report") -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    pdf.cell(0, 8, title.encode("latin-1", errors="replace").decode("latin-1"), ln=True)
    pdf.ln(2)
    for line in content.splitlines():
        text = line.encode("latin-1", errors="replace").decode("latin-1") or " "
        pdf.cell(0, 5, text[:200], ln=True)
    out = pdf.output()
    if isinstance(out, bytes):
        return out
    if isinstance(out, bytearray):
        return bytes(out)
    return str(out).encode("latin-1")


def export_docx(content: str, *, title: str = "ExposureFlow Report") -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)
    for line in content.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line.strip())
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_report(content: str, fmt: ExportFormat, *, title: str = "ExposureFlow Report") -> bytes:
    if fmt == "markdown":
        return export_markdown(content)
    if fmt == "pdf":
        return export_pdf(content, title=title)
    if fmt == "docx":
        return export_docx(content, title=title)
    raise ValueError(f"Unsupported format: {fmt}")
